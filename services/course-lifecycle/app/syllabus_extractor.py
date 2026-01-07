
import os
import pdfplumber
import fitz # PyMuPDF
import docx
import json
import google.generativeai as genai
from typing import Dict, Any, Optional
import re

# Configure Gemini (Lazy init in function)
# GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
# if GEMINI_API_KEY:
#     genai.configure(api_key=GEMINI_API_KEY)

BLUEPRINT_SCHEMA = """
{
  "course_identity": {
    "course_code": "string",
    "course_name": "string",
    "semester": "string (optional)",
    "credits": "string (optional)",
    "ltp": "string (L-T-P format, e.g. 3-0-0)",
    "prerequisites": ["string (list of prerequisite courses)"]
  },
  "description": "string (course description)",
  "course_objectives": ["string (verbatim from syllabus)"],
  "course_outcomes": [
      {"id": "string (e.g. CO1)", "description": "string (verbatim outcome text)"}
  ],
  "modules": [
    {
      "id": "string (e.g. UNIT 1)",
	    "name": "string (Module/Unit Title)",
	    "title": "string (Deprecated alias for name; if present copy into name)",
	    "duration": "number (minutes; if missing use 600)",
      "module_outcome": "string (Specific module outcome if present, else 'Not Provided')",
      "topics": [
          {
              "id": "string (e.g. U1T1)",
              "name": "string (Topic Name)",
              "topic_outcome": "string (default 'Not Provided')"
          }
      ]
    }
  ],
  "textbooks": [
      {"citation": "string (Full citation)"}
  ],
  "reference_books": [
      {"citation": "string (Full citation)"}
  ]
}
"""


def _parse_course_outcomes_from_text(text: str) -> Dict[str, str]:
    """Extract CO descriptions from raw syllabus text (best effort).

    Supports common patterns like:
      - CO1: ...
      - CO 1 - ...
      - CO-2: ...
    """

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    out: Dict[str, str] = {}
    i = 0
    pat = re.compile(r"\bCO\s*[-:]?\s*(\d+)\s*[:\-]\s*(.+)$", re.IGNORECASE)

    while i < len(lines):
        m = pat.search(lines[i])
        if not m:
            i += 1
            continue

        co_id = f"CO{m.group(1)}"
        desc_parts = [m.group(2).strip()]

        # Stitch continuation lines until next CO-like line
        j = i + 1
        while j < len(lines):
            if pat.search(lines[j]):
                break
            # Stop if we hit a new section header-ish line
            if re.match(r"^(UNIT|MODULE|UNIT\s+\w+|MODULE\s+\w+|TEXTBOOKS|REFERENCES|REFERENCE BOOKS)\b", lines[j], re.IGNORECASE):
                break
            # Append short continuation lines
            desc_parts.append(lines[j])
            j += 1

        desc = " ".join([p for p in desc_parts if p]).strip()
        if desc:
            out[co_id.upper()] = desc

        i = j

    return out


def _normalize_blueprint(bp: Dict[str, Any], raw_text: Optional[str] = None) -> Dict[str, Any]:
    """Normalize blueprint keys/values so the frontend doesn't render blanks.

    - Ensures module.name exists (copies from title if needed)
    - Ensures module.duration exists and is a positive integer (defaults to 600)
    - Ensures course_outcomes[i].description is not empty (fills from raw text or 'Not Provided')
    """

    if not isinstance(bp, dict):
        return bp

    co_from_text: Dict[str, str] = {}
    if raw_text:
        try:
            co_from_text = _parse_course_outcomes_from_text(raw_text)
        except Exception:
            co_from_text = {}

    # Course outcomes
    cos = bp.get("course_outcomes")
    if isinstance(cos, list):
        for co in cos:
            if not isinstance(co, dict):
                continue
            co_id = (co.get("id") or "").strip()
            desc = (co.get("description") or "").strip()
            if not desc:
                desc = co_from_text.get(co_id.upper(), "Not Provided")
            co["id"] = co_id or co.get("id") or "CO?"
            co["description"] = desc

    # Modules
    mods = bp.get("modules")
    if isinstance(mods, list):
        for idx, m in enumerate(mods):
            if not isinstance(m, dict):
                continue
            name = (m.get("name") or "").strip()
            if not name:
                name = (m.get("title") or "").strip() or f"Module {idx+1}"
            m["name"] = name
            # Duration in minutes
            dur = m.get("duration")
            try:
                dur_int = int(dur) if dur is not None else 0
            except Exception:
                dur_int = 0
            if dur_int <= 0:
                dur_int = 600
            m["duration"] = dur_int

            # Ensure outcomes are not empty strings
            if not (m.get("module_outcome") or "").strip():
                m["module_outcome"] = "Not Provided"

            topics = m.get("topics")
            if isinstance(topics, list):
                for t in topics:
                    if not isinstance(t, dict):
                        continue
                    if not (t.get("name") or "").strip() and (t.get("title") or "").strip():
                        t["name"] = (t.get("title") or "").strip()
                    if not (t.get("topic_outcome") or "").strip():
                        t["topic_outcome"] = "Not Provided"

    return bp

from pathlib import Path

async def extract_text_from_file(file, filename: str) -> str:
    content = ""
    filename = filename.lower()
    
    if filename.endswith(".pdf"):
        # Check settings for OCR
        from .settings import settings
        use_ocr = settings.ENABLE_OCR
        
        # Determine bytes or path
        file_bytes = None
        file_path = None
        
        if isinstance(file, (str, Path)):
            file_path = file
            try:
                with open(file, "rb") as f:
                    file_bytes = f.read()
            except Exception:
                pass
        else:
            # Stream / UploadFile
            try:
                file_bytes = await file.read()
                await file.seek(0) # Reset
            except Exception:
                pass

        if use_ocr and file_bytes:
            from .ocr_utils import extract_text_from_pdf as ocr_extract
            try:
                logger.info(f"Attempting OCR for {filename}")
                content = ocr_extract(file_bytes, filename)
            except Exception as e:
                logger.error(f"OCR failed: {e}")
                
        # Fallback to pdfplumber if no content via OCR
        if not content:
            try:
                if file_path:
                     with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            content += page.extract_text() or ""
                else:
                    # Stream logic for pdfplumber can be tricky with async file
                    # Better to use bytes io if possible
                    import io
                    if file_bytes:
                         with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                            for page in pdf.pages:
                                content += page.extract_text() or ""
                    else:
                         # Try direct stream if supported
                         with pdfplumber.open(file) as pdf:
                            for page in pdf.pages:
                                content += page.extract_text() or ""
            except Exception as e:
                logger.error(f"PDFPlumber failed: {e}")
            
    elif filename.endswith(".docx"):
        # Handle Docx
        try:
            if isinstance(file, str):
                 doc = docx.Document(file)
            else:
                 # Ensure proper stream handling for docx
                 import io
                 b = await file.read()
                 await file.seek(0)
                 doc = docx.Document(io.BytesIO(b))
                 
            content = "\\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
             logger.error(f"Reading DOCX failed: {e}")
             
    else:
        # Assume text/markdown
        try:
            if isinstance(file, str):
                with open(file, 'r', encoding='utf-8', errors='ignore') as f: content = f.read()
            else:
                content = (await file.read()).decode("utf-8", errors="ignore")
                await file.seek(0) # Reset pointer
        except Exception as e:
             logger.error(f"Reading Text failed: {e}")
            
    return content

async def generate_blueprint_from_text(text: str, api_key: Optional[str] = None) -> Dict[str, Any]:
    if not api_key:
        # Fallback to env if not passed, but ideally passed
        api_key = os.getenv("GEMINI_API_KEY")
        
    if not api_key:
        raise Exception("GEMINI_API_KEY not set")

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('models/gemini-2.0-flash-exp') # Upgraded model for better extraction
    
    prompt = f"""
    You are a Strict Data Extraction Engine. Extract a structured 'Course Blueprint' from the syllabus text below.
    
	    STRICT RULES:
	    1. EXTRACT DATA VERBATIM. Do not summarize or rewrite outcomes.
	    2. DO NOT INFER OR INVENT. If a field is missing, set it to "Not Provided" or empty list [].
	       - NEVER return empty strings. Use "Not Provided".
    3. HIERARCHY: Structure specifically as Modules (Units) -> Topics. 
       - Split topics by commas, semi-colons, or newlines in the syllabus. 
       - Ensure every topic listed in the syllabus corresponds to a topic object.
    4. MAPPING: Map 'UNIT X' headings to 'modules'.
	    5. DURATION: If module duration is not explicitly stated, set duration = 600 (minutes).
    
    Output JSON STRICTLY adhering to this schema:
    {BLUEPRINT_SCHEMA}
    
    Syllabus Text:
    {text[:40000]} 
    """

    response = model.generate_content(
        prompt,
        generation_config={"response_mime_type": "application/json"},
    )
    try:
        bp = json.loads(response.text)
        return _normalize_blueprint(bp, raw_text=text)
    except Exception as e:
        print(f"JSON Parse Error: {response.text}")
        return {"error": "Failed to parse blueprint", "raw": response.text}
